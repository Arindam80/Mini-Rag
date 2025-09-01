



import os
import io
import tempfile
from dotenv import load_dotenv
from typing import Optional

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from qdrant_client import QdrantClient

# Document Processing Imports
import PyPDF2
import docx
import mammoth
import traceback

# LangChain Imports
from langchain_cohere import CohereEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain.prompts import ChatPromptTemplate
from langchain_cohere import CohereRerank
from langchain_groq import ChatGroq
from cohere import Client
from langchain.retrievers import ContextualCompressionRetriever
from langchain_community.vectorstores import Qdrant

# Import specific Cohere error classes
from cohere.errors import (
    BadRequestError,
    UnauthorizedError,
    TooManyRequestsError,
    InternalServerError,
    ServiceUnavailableError,
    NotFoundError,
    ForbiddenError,
    UnprocessableEntityError
)

# Create a tuple for catching multiple Cohere errors
COHERE_ERRORS = (
    BadRequestError,
    UnauthorizedError,
    TooManyRequestsError,
    InternalServerError,
    ServiceUnavailableError,
    NotFoundError,
    ForbiddenError,
    UnprocessableEntityError
)

# --- INITIAL SETUP ---

# Load environment variables from .env file
load_dotenv()

# Instantiate the FastAPI app
app = FastAPI(title="AI Resume Analyzer", description="Enhanced RAG system with resume analysis capabilities")

# --- CORS MIDDLEWARE ---
origins = [
    "http://localhost:5500",
    "http://127.0.0.1:5500",
    # "https://mini-rag-sepia.vercel.app",
    "http://localhost:3000",
    "http://127.0.0.1:3000"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- GLOBAL VARIABLES & CONSTANTS ---
QDRANT_COLLECTION_NAME = "my_rag_collection_cohere"
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}

# --- PYDANTIC MODELS ---
class UploadData(BaseModel):
    text: str

class QueryData(BaseModel):
    question: str

# --- DOCUMENT PROCESSING FUNCTIONS ---

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")

def extract_text_from_doc(file_content: bytes) -> str:
    """Extract text from DOC file using mammoth."""
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_file.flush()
            
            with open(tmp_file.name, 'rb') as f:
                result = mammoth.extract_raw_text(f)
                text = result.value
            
            os.unlink(tmp_file.name)  # Clean up temp file
            return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOC: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file."""
    try:
        return file_content.decode('utf-8').strip()
    except UnicodeDecodeError:
        try:
            return file_content.decode('latin-1').strip()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading TXT file: {str(e)}")

def process_file(file_content: bytes, filename: str) -> str:
    """Process uploaded file and extract text based on file extension."""
    file_ext = os.path.splitext(filename)[1].lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_content)
    elif file_ext == '.docx':
        return extract_text_from_docx(file_content)
    elif file_ext == '.doc':
        return extract_text_from_doc(file_content)
    elif file_ext == '.txt':
        return extract_text_from_txt(file_content)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")

def chunk_and_embed_text(text: str) -> int:
    """Chunk text and store embeddings in Qdrant."""
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, 
            chunk_overlap=150,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        chunks = text_splitter.split_text(text)
        
        # Create documents with metadata
        documents = [
            Document(
                page_content=chunk, 
                metadata={
                    "position": i + 1,
                    "chunk_length": len(chunk),
                    "source_type": "uploaded_document"
                }
            ) 
            for i, chunk in enumerate(chunks)
        ]
        
        embeddings = CohereEmbeddings(
            cohere_api_key=os.getenv("COHERE_API_KEY"), 
            model="embed-english-v3.0"
        )
        
        QdrantVectorStore.from_documents(
            documents,
            embeddings,
            url=os.getenv("QDRANT_URL"),
            api_key=os.getenv("QDRANT_API_KEY"),
            collection_name=QDRANT_COLLECTION_NAME,
            force_recreate=True,
        )
        
        return len(documents)
    except COHERE_ERRORS as e:
        raise HTTPException(status_code=500, detail=f"Cohere API error: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing text: {str(e)}")

# --- API ENDPOINTS ---

@app.get("/")
def read_root():
    """Health check endpoint."""
    return {
        "status": "AI Resume Analyzer API is running",
        "version": "2.0",
        "features": ["Text upload", "Resume file upload", "AI analysis"]
    }

@app.post("/upload")
async def upload(data: UploadData):
    """
    Endpoint for uploading, chunking, embedding, and storing text in the vector database.
    """
    try:
        if not data.text.strip():
            raise HTTPException(status_code=400, detail="Text content cannot be empty")
        
        chunk_count = chunk_and_embed_text(data.text)
        
        return {
            "message": f"Successfully processed and uploaded {chunk_count} chunks",
            "chunks": chunk_count,
            "status": "success"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """
    Endpoint for uploading resume files (PDF, DOC, DOCX, TXT).
    """
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file uploaded")
        
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
            )
        
        # Read file content
        file_content = await file.read()
        
        if len(file_content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")
        
        if len(file_content) == 0:
            raise HTTPException(status_code=400, detail="File is empty")
        
        # Extract text from file
        extracted_text = process_file(file_content, file.filename)
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No text content found in the file")
        
        # Process and store in vector database
        chunk_count = chunk_and_embed_text(extracted_text)
        
        return {
            "message": f"Successfully analyzed resume '{file.filename}' and created {chunk_count} chunks",
            "filename": file.filename,
            "chunks": chunk_count,
            "text_length": len(extracted_text),
            "status": "success"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Resume processing failed: {str(e)}")


@app.post("/query")
async def query(data: QueryData):
    """
    Endpoint for querying the RAG pipeline with enhanced resume-specific prompts.
    """
    try:
        if not data.question.strip():
            raise HTTPException(status_code=400, detail="Question cannot be empty")
        
        # Initialize components
        embeddings = CohereEmbeddings(
            cohere_api_key=os.getenv("COHERE_API_KEY"), 
            model="embed-english-v3.0"
        )
        llm = ChatGroq(
            groq_api_key=os.getenv("GROQ_API_KEY"), 
            model_name="llama-3.1-8b-instant", 
            temperature=0
        )
        
        client = QdrantClient(
            url=os.getenv("QDRANT_URL"),
            api_key=os.getenv("QDRANT_API_KEY"),
        )

        # ✅ FIX: include embeddings
        vector_store = Qdrant(
            client=client,
            collection_name=QDRANT_COLLECTION_NAME,
            embeddings=embeddings,
        )

        # Reranker
        reranker = CohereRerank(
            cohere_api_key=os.getenv("COHERE_API_KEY"),
            model="rerank-english-v3.0",
            top_n=5
        )

        # Retriever
        retriever = vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 10}
        )

        # ✅ FIX: use compress_documents instead of invoke()
        # retrieved_docs = retriever.get_relevant_documents(data.question)
        retrieved_docs = retriever.invoke(data.question)

        reranked_docs = reranker.compress_documents(
            documents=retrieved_docs, 
            query=data.question
        )

        if not reranked_docs:
            raise HTTPException(
                status_code=404, 
                detail="No relevant content found. Please upload a document first."
            )
        
        # Build context
        context_text = "\n\n".join([
            f"[{doc.metadata['position']}] {doc.page_content}" 
            for doc in reranked_docs
        ])
        
        # Prompt
        prompt_template = """
        You are an expert HR consultant and resume analyzer...
        Context:
        {context}
        Question: {question}
        Answer:
        """
        prompt = ChatPromptTemplate.from_template(prompt_template)
        chain = prompt | llm
        response_llm = chain.invoke({
            "context": context_text, 
            "question": data.question
        })
        
        # Sources
        source_documents = [
            {
                "content": doc.page_content[:500] + "..." if len(doc.page_content) > 500 else doc.page_content,
                "position": doc.metadata['position']
            } 
            for doc in reranked_docs
        ]
        
        return {
            "answer": response_llm.content,
            "sources": source_documents,
            "question": data.question,
            "status": "success"
        }

    except COHERE_ERRORS as e:
        raise HTTPException(status_code=500, detail=f"Cohere API error: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        print("❌ Query failed with exception:")
        traceback.print_exc()   # <--- this will show the real error in your terminal
        raise HTTPException(status_code=500, detail=f"Query processing failed: {str(e)}")
@app.get("/health")
def health_check():
    """Detailed health check endpoint."""
    try:
        # Test environment variables
        required_vars = ["COHERE_API_KEY", "GROQ_API_KEY", "QDRANT_URL", "QDRANT_API_KEY"]
        missing_vars = [var for var in required_vars if not os.getenv(var)]
        
        if missing_vars:
            return {
                "status": "unhealthy",
                "error": f"Missing environment variables: {', '.join(missing_vars)}"
            }
        
        return {
            "status": "healthy",
            "database": "connected",
            "ai_services": "available",
            "file_upload": "enabled",
            "max_file_size": f"{MAX_FILE_SIZE // 1024 // 1024}MB",
            "supported_formats": list(ALLOWED_EXTENSIONS)
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e)
        }